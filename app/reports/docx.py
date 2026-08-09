from __future__ import annotations

import uuid
from datetime import UTC, datetime
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.core.config import settings
from app.reports.visuals import weather_snapshot
from app.storage.database import save_report

BLACK = RGBColor(0, 0, 0)
LIGHT_GRAY = "E7E7E7"
VERY_LIGHT_GRAY = "F5F5F5"
FONT_NAME = "Times New Roman"


def _label(value: Any) -> str:
    return str(value or "").replace("_", " ").strip().title()


def _number(value: Any, digits: int = 2) -> str:
    if value is None:
        return "Not available"
    try:
        number = float(value)
    except (TypeError, ValueError):
        return str(value)
    if not (number == number and abs(number) != float("inf")):
        return "Not available"
    if abs(number - round(number)) < 1e-10:
        return f"{int(round(number)):,}"
    return f"{number:,.{digits}f}".rstrip("0").rstrip(".")


def _percent(value: Any, digits: int = 1, already_percent: bool = False) -> str:
    if value is None:
        return "Not available"
    try:
        number = float(value)
    except (TypeError, ValueError):
        return str(value)
    if not already_percent:
        number *= 100.0
    return f"{number:,.{digits}f}%"


def _text(value: Any, max_chars: int = 1400) -> str:
    if value is None:
        return "Not available"
    if isinstance(value, bool):
        return "Yes" if value else "No"
    if isinstance(value, float):
        return _number(value, 4)
    if isinstance(value, (list, tuple)):
        result = "; ".join(_text(item, 180) for item in value[:30])
    elif isinstance(value, dict):
        result = "; ".join(f"{_label(key)}: {_text(item, 180)}" for key, item in list(value.items())[:30])
    else:
        result = str(value)
    result = result.replace("–", "-").replace("—", "-")
    return result if len(result) <= max_chars else result[: max_chars - 3] + "..."


def _set_run_font(run, *, size: float | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.color.rgb = BLACK
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top: int = 80, start: int = 90, bottom: int = 80, end: int = 90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _format_document(document: Document) -> None:
    styles = document.styles
    for style in styles:
        if hasattr(style, "font"):
            style.font.name = FONT_NAME
            style.font.color.rgb = BLACK
            try:
                style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
            except AttributeError:
                pass
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.size = Pt(18)
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.size = Pt(12)
    styles["Heading 2"].font.bold = True
    styles["Heading 3"].font.size = Pt(11)
    styles["Heading 3"].font.bold = True


def _paragraph(document: Document, text: Any = "", *, bold: bool = False, italic: bool = False,
               align: WD_ALIGN_PARAGRAPH | None = None, size: float = 10.5,
               space_after: float = 4.0) -> None:
    paragraph = document.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.add_run(_text(text))
    _set_run_font(run, size=size, bold=bold, italic=italic)


def _heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    paragraph.paragraph_format.space_after = Pt(4)
    for run in paragraph.runs:
        _set_run_font(run, size={1: 14, 2: 12, 3: 11}.get(level, 10.5), bold=True)


def _style_table(table, header: bool = True) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        if header and row_index == 0:
            tbl_header = OxmlElement("w:tblHeader")
            tbl_header.set(qn("w:val"), "true")
            tr_pr.append(tbl_header)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)
            if header and row_index == 0:
                _set_cell_shading(cell, LIGHT_GRAY)
            elif row_index % 2 == 0:
                _set_cell_shading(cell, VERY_LIGHT_GRAY)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    _set_run_font(run, size=8.5, bold=(header and row_index == 0))


def _key_value_table(document: Document, values: Iterable[tuple[str, Any]] | dict[str, Any]) -> None:
    items = list(values.items()) if isinstance(values, dict) else list(values)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Value"
    for key, value in items:
        cells = table.add_row().cells
        cells[0].text = _label(key)
        cells[1].text = _text(value)
    _style_table(table)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _bullet_list(document: Document, items: Iterable[Any]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(_text(item))
        _set_run_font(run, size=10)


def _annual_product_table(document: Document, annual_rows: list[dict[str, Any]]) -> None:
    if not annual_rows:
        _paragraph(document, "No annual three-product outlook was attached to this report.")
        return
    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headings = ["Year", "Coconut w/ husk (t)", "Mature (t)", "Young (t)", "Mature share", "Young share"]
    for cell, value in zip(table.rows[0].cells, headings):
        cell.text = value
    indexes = sorted(set([0, len(annual_rows) - 1] + list(range(4, len(annual_rows), 5))))
    for index in indexes:
        item = annual_rows[index]
        cells = table.add_row().cells
        values = [
            item.get("year"),
            _number(item.get("coconut_w_husk_tons"), 3),
            _number(item.get("coconut_mature_tons"), 3),
            _number(item.get("coconut_young_tons"), 3),
            _percent(item.get("mature_share")),
            _percent(item.get("young_share")),
        ]
        for cell, value in zip(cells, values):
            cell.text = _text(value)
    _style_table(table)


def _event_table(document: Document, events: list[dict[str, Any]]) -> None:
    if not events:
        _paragraph(document, "No major event periods were flagged in the selected trajectory.")
        return
    table = document.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headings = ["Event", "Start", "End", "Duration", "Severity", "Estimated loss (t)", "Trees for inspection"]
    for cell, value in zip(table.rows[0].cells, headings):
        cell.text = value
    for event in events[:35]:
        severity = event.get("severity_percent")
        if severity is None:
            severity = float(event.get("peak_severity", 0.0)) * 100.0
        values = [
            event.get("label"), event.get("start_date"), event.get("end_date"),
            f"{_number(event.get('duration_weeks', event.get('weeks')), 0)} weeks", f"{_number(severity, 1)}/100",
            _number(event.get("estimated_production_loss_tons"), 3),
            _number(event.get("estimated_trees_affected"), 0),
        ]
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = _text(value)
    _style_table(table)


def _pest_table(document: Document, pests: list[dict[str, Any]]) -> None:
    if not pests:
        _paragraph(document, "No pest-specific assessment was attached.")
        return
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headings = ["Pest", "Scientific name", "Score", "Risk class", "Primary affected part"]
    for cell, value in zip(table.rows[0].cells, headings):
        cell.text = value
    for pest in pests:
        values = [pest.get("common_name"), pest.get("scientific_name"), f"{_number(pest.get('outbreak_score'), 1)}/100", pest.get("risk_class"), pest.get("affected_part")]
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = _text(value)
    _style_table(table)


def _add_snapshot(document: Document, frame: dict[str, Any], farm_position: dict[str, Any], index: int) -> None:
    marker = {"map_x": farm_position.get("x", 0.5), "map_y": farm_position.get("y", 0.5)}
    image_data = weather_snapshot(frame, marker, width=1180, height=640)
    title = frame.get("label") or frame.get("week_start") or frame.get("date") or f"Critical date {index}"
    _heading(document, f"Figure {index}. Critical weather frame: {title}", level=2)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(image_data, width=Inches(6.65))
    caption = (
        f"Rainfall {_number(frame.get('rainfall_mm'), 1)} mm; peak intensity "
        f"{_number(frame.get('rain_intensity_mm_h'), 1)} mm/h; maximum temperature "
        f"{_number(frame.get('temperature_max_c', frame.get('temperature_c')), 1)} C; "
        f"farm condition {_percent(frame.get('farm_condition_score'))}. The colored surface is the "
        "forecast or climate-conditioned rain field used by the farm simulation."
    )
    _paragraph(document, caption, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)


def _footer(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("COCO-AID Decision-Support Research Report")
    _set_run_font(run, size=8)


def generate_docx(analysis: dict[str, Any], analysis_id: str | None = None) -> tuple[str, Path]:
    if "result" in analysis and isinstance(analysis["result"], dict):
        analysis = analysis["result"]
    settings.reports_dir.mkdir(parents=True, exist_ok=True)
    report_id = str(uuid.uuid4())
    path = settings.reports_dir / f"COCO-AID_Report_{report_id[:8]}.docx"

    document = Document()
    _format_document(document)
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    _footer(section)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(5)
    run = title.add_run("COCO-AID DECISION-SUPPORT REPORT")
    _set_run_font(run, size=18, bold=True)
    subtitle = (
        "Bayesian Probabilistic Agroecosystem Simulation and Geospatial AI-Based "
        "Decision-Support Framework for Coconut Rehabilitation"
    )
    _paragraph(document, subtitle, align=WD_ALIGN_PARAGRAPH.CENTER, size=10, space_after=9)
    warning = analysis.get("scientific_warning") or (
        "Official historical records are used where available. Missing farm measurements and future "
        "conditions are model-estimated and require field validation."
    )
    _paragraph(document, warning, bold=True, size=10, space_after=8)

    overview = analysis.get("overview") or {}
    farm = analysis.get("farm_assessment") or {}
    official = analysis.get("official_production_reference") or {}
    climate = analysis.get("climate_projection") or {}
    simulation = analysis.get("recommended_simulation") or {}
    comparison = analysis.get("scenario_comparison") or {}
    pest = analysis.get("pest_risk") or {}
    pest_specific = analysis.get("pest_specific") or {}
    suitability = analysis.get("land_suitability") or {}
    rehab = analysis.get("rehabilitation_map") or {}
    rehab_events = analysis.get("rehabilitation_event_plans") or {}
    forecast = analysis.get("farm_site_forecast") or {}
    health_snapshot = analysis.get("farm_health_snapshot") or {}
    metadata = analysis.get("metadata") or {}

    _heading(document, "1. Executive Summary")
    _key_value_table(document, [
        ("Farm", overview.get("farm_name", forecast.get("farm", {}).get("name"))),
        ("Current annual production", f"{_number(overview.get('current_production_tons'), 3)} metric tons"),
        ("Projected end-year median", f"{_number(overview.get('projected_end_median_tons'), 3)} metric tons"),
        ("Rehabilitation probability", _percent(overview.get("rehabilitation_probability"))),
        ("Severe-loss probability", _percent(overview.get("severe_loss_probability"))),
        ("Bayesian pest posterior", _percent(overview.get("pest_risk_probability", pest.get("posterior_probability")))),
        ("Land suitability", _percent(overview.get("land_suitability_percentage", suitability.get("percentage")), already_percent=True)),
        ("Recommended intervention", _label(overview.get("recommended_intervention", comparison.get("recommended_intervention")))),
        ("Climate scenario", overview.get("climate_scenario", forecast.get("scenario"))),
    ])

    _heading(document, "2. Farm Profile and Data Quality")
    _key_value_table(document, [
        ("Entered area", f"{_number(farm.get('entered_area_hectares'), 3)} hectares"),
        ("Polygon area", f"{_number(farm.get('polygon_area_hectares'), 3)} hectares" if farm.get("polygon_area_hectares") is not None else "Not provided"),
        ("Tree density", f"{_number(farm.get('tree_density_per_hectare'), 2)} trees per hectare"),
        ("Productive and recovering fraction", _percent(farm.get("productive_recovering_fraction"))),
        ("At-risk tree fraction", _percent(farm.get("at_risk_fraction"))),
        ("Calculated yield", f"{_number(farm.get('calculated_yield_tons_per_hectare'), 3)} t/ha"),
        ("Data quality score", _percent(farm.get("data_quality_score"))),
        ("Data quality class", farm.get("data_quality_class")),
    ])
    _bullet_list(document, farm.get("warnings") or [])

    _heading(document, "3. Official Provincial Production Reference")
    products = official.get("products") or {}
    _key_value_table(document, [
        ("Province", official.get("province")),
        ("Region", official.get("region")),
        ("Reference level", official.get("reference_level")),
        ("Latest official Coconut w/ husk", f"{_number(products.get('coconut_w_husk', {}).get('latest_official_2025_tons'), 2)} t"),
        ("Latest official Coconut Mature", f"{_number(products.get('coconut_mature', {}).get('latest_official_2025_tons'), 2)} t"),
        ("Latest official Coconut Young", f"{_number(products.get('coconut_young', {}).get('latest_official_2025_tons'), 2)} t"),
        ("Source", official.get("metadata", {}).get("source")),
        ("Table code", official.get("metadata", {}).get("table_code")),
    ])

    _heading(document, "4. Climate and Farm Production Outlook")
    annual_summary = climate.get("annual_summary") or {}
    _key_value_table(document, [
        ("Climate period", climate.get("display_label", climate.get("period"))),
        ("Scenario", climate.get("scenario", forecast.get("scenario"))),
        ("Mean temperature", f"{_number(annual_summary.get('mean_temperature_c'), 2)} C"),
        ("Annual rainfall", f"{_number(annual_summary.get('annual_precipitation_mm'), 1)} mm"),
        ("Forecast horizon", f"{forecast.get('effective_start_date', 'Not available')} to {forecast.get('effective_end_date', 'Not available')}"),
        ("Timeline resolution", forecast.get("timeline_resolution")),
        ("Monte Carlo runs", forecast.get("runs", metadata.get("simulation_count_per_intervention"))),
    ])
    _paragraph(document, "Long-term weekly weather fields are climate-conditioned scenario paths, not exact forecasts of future clouds, rainfall, heat waves, or storm dates.")

    _heading(document, "5. Critical Weather Dates")
    critical = forecast.get("critical_weather_frames") or []
    if critical:
        farm_position = forecast.get("farm_map_position") or {}
        for index, frame in enumerate(critical[:6], start=1):
            _add_snapshot(document, frame, farm_position, index)
    else:
        _paragraph(document, "No critical-date weather snapshots were attached to this report.")

    document.add_page_break()
    _heading(document, "6. Three-Product Production Projection")
    _annual_product_table(document, forecast.get("annual_by_product") or [])
    product_model = forecast.get("product_model") or {}
    if product_model:
        _heading(document, "Weather-responsive product model", level=2)
        _key_value_table(document, product_model)

    _heading(document, "7. Extreme-Weather Risk Timeline")
    _event_table(document, forecast.get("extreme_events") or [])
    _paragraph(document, "Estimated loss is calculated from event type, peak and mean severity, event duration, and baseline weekly production. Consequently, reported loss increases consistently as event severity or duration increases.")

    _heading(document, "8. Bayesian and Pest-Specific Risk Assessment")
    _key_value_table(document, [
        ("Bayesian prior", _percent(pest.get("prior_probability"))),
        ("Bayesian posterior", _percent(pest.get("posterior_probability"))),
        ("Bayesian risk class", pest.get("risk_class")),
        ("Highest pest-specific outbreak score", f"{_number(pest_specific.get('highest_outbreak_score'), 1)}/100"),
        ("Top pest-specific risk", pest_specific.get("top_risk_pest")),
        ("Overall pest pressure", f"{_number(pest_specific.get('overall_outbreak_pressure'), 1)}/100"),
    ])
    pest_rows = pest_specific.get("pests") or []
    _pest_table(document, pest_rows)
    for pest_row in pest_rows[:5]:
        _heading(document, f"{pest_row.get('common_name', 'Pest')} recommendations", level=2)
        _bullet_list(document, pest_row.get("ai_recommendations") or [])

    _heading(document, "9. Land Suitability and Farm Health")
    _key_value_table(document, [
        ("Suitability percentage", _percent(suitability.get("percentage"), already_percent=True)),
        ("Suitability class", suitability.get("class")),
        ("Limiting factors", suitability.get("limiting_factors")),
        ("Farm condition", _percent(health_snapshot.get("farm_condition_score"))),
        ("Baseline current-condition grid", f"{rehab.get('rows', 'Not available')} x {rehab.get('cols', 'Not available')}"),
        ("Baseline high-priority cells", health_snapshot.get("rehabilitation_summary", {}).get("high_priority_cells")),
        ("Baseline visible cells", len(rehab.get("cells") or [])),
    ])
    components = suitability.get("component_scores") or {}
    if components:
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Suitability component"
        table.rows[0].cells[1].text = "Membership score"
        for key, value in components.items():
            cells = table.add_row().cells
            cells[0].text = _label(key)
            cells[1].text = _percent(value)
        _style_table(table)
    plans = rehab_events.get("plans") or []
    if plans:
        _heading(document, "Event-Linked Rehabilitation Schedule", level=2)
        table = document.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        for cell, value in zip(table.rows[0].cells, ["Event", "Event period", "Inspect", "Rehabilitation", "Yellow / Red"]):
            cell.text = value
        for plan in plans[:12]:
            values = [
                plan.get("event_label"),
                f"{plan.get('event_start_date')} to {plan.get('event_end_date')}",
                plan.get("recommended_assessment_date"),
                plan.get("recommended_rehabilitation_date"),
                f"{(plan.get('counts') or {}).get('Needs inspection', 0)} / {(plan.get('counts') or {}).get('Needs Rehabilitation', 0)}",
            ]
            cells = table.add_row().cells
            for cell, value in zip(cells, values):
                cell.text = _text(value)
        _style_table(table)
        _paragraph(document, "Green zones indicate no immediate damage action, yellow zones require field inspection, and red zones indicate likely rehabilitation after field verification.")

    _heading(document, "10. Intervention Comparison")
    ranking = comparison.get("ranking") or []
    if ranking:
        table = document.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        headings = ["Rank", "Intervention", "Utility", "Final median", "Recovery", "Severe loss"]
        for cell, value in zip(table.rows[0].cells, headings):
            cell.text = value
        for item in ranking:
            values = [
                item.get("rank"), _label(item.get("intervention")), _number(item.get("expected_utility"), 3),
                f"{_number(item.get('final_median_tons'), 3)} t", _percent(item.get("rehabilitation_probability")),
                _percent(item.get("severe_loss_probability")),
            ]
            cells = table.add_row().cells
            for cell, value in zip(cells, values):
                cell.text = _text(value)
        _style_table(table)
    else:
        _paragraph(document, "No intervention comparison was available.")

    _heading(document, "Model Versions, Provenance, and Limitations", level=2)
    _key_value_table(document, [
        ("Calculation version", metadata.get("calculation_version")),
        ("Parameter version", metadata.get("parameter_version")),
        ("Model versions", metadata.get("model_versions")),
        ("Random seed", metadata.get("random_seed", forecast.get("seed"))),
        ("Generated at", datetime.now(UTC).isoformat()),
        ("Data source type", metadata.get("data_source_type", forecast.get("data_source_type"))),
    ])
    limitations: list[str] = []
    for source in [metadata, simulation, comparison, forecast, pest_specific]:
        if isinstance(source, dict) and isinstance(source.get("limitations"), list):
            limitations.extend(str(item) for item in source["limitations"])
    limitations.extend([
        "The long-term weather path is a plausible climate-conditioned scenario, not an exact forecast to 2050.",
        "Pest-specific scores are inspection priorities and do not replace laboratory or expert identification.",
        "Farm-scale accuracy requires locally measured soil, pest, weather, tree-state, and production records.",
    ])
    _bullet_list(document, list(dict.fromkeys(limitations)))

    document.save(path)
    save_report(report_id, path, analysis_id, report_type="docx")
    return report_id, path
