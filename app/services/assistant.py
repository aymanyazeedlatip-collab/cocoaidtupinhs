from __future__ import annotations

import json
import re
import uuid
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

import httpx
from docx import Document
from pypdf import PdfReader

from app.core.config import settings
from app.storage.database import get_report, report_record
from app.services.supabase_state import supabase_state

_PRIVATE_FILE = settings.private_settings_path
_DOC_DIR = settings.cache_dir / "assistant_documents"
_MAX_EXTRACTED_CHARS = 120_000
_MODEL_FALLBACKS = ("gemini-flash-latest", "gemini-2.5-flash", "gemini-2.5-flash-lite")

SYSTEM_PROMPT = """You are CoCO-PILOT, the coconut-farming assistant inside COCO-AID.
Give concise, practical, evidence-aware advice for coconut farms in the Philippines.
Use the supplied COCO-AID farm, forecast, pest, suitability, hazard, and report context when present.
Never present a climate-conditioned long-term projection as an exact weather forecast.
Never invent values that are missing. State uncertainty plainly.
Prefer integrated pest management, field inspection, sanitation, monitoring, biological control, and local agricultural-extension consultation.
Do not prescribe unverified pesticide doses; direct the farmer to the product label and local agriculture office.
Format replies for a compact chat window: a short conclusion, then at most four bullets, and a final action line. Avoid long paragraphs.
When discussing percentages, explain what the percentage means and what action threshold it suggests.
"""


def _read_private() -> dict[str, Any]:
    try:
        return json.loads(_PRIVATE_FILE.read_text(encoding="utf-8"))
    except (FileNotFoundError, json.JSONDecodeError, OSError):
        return {}


def assistant_status() -> dict[str, Any]:
    private = _read_private()
    configured = bool(settings.gemini_api_key or private.get("gemini_api_key"))
    storage = "Environment variable" if settings.gemini_api_key else ("Local private settings file" if configured else "Not configured")
    resolved = private.get("resolved_model")
    return {
        "configured": configured,
        "provider": "Google AI",
        "model": "Automatic compatible Flash model",
        "resolved": bool(resolved),
        "storage": storage,
    }


def save_api_key(api_key: str) -> None:
    value = api_key.strip()
    if not value or len(value) < 20 or len(value) > 300:
        raise ValueError("Enter a valid Gemini API key.")
    _PRIVATE_FILE.parent.mkdir(parents=True, exist_ok=True)
    existing = _read_private()
    payload = {
        "gemini_api_key": value,
        "updated_at": datetime.now(UTC).isoformat(),
    }
    _PRIVATE_FILE.write_text(json.dumps(payload, indent=2), encoding="utf-8")
    try:
        _PRIVATE_FILE.chmod(0o600)
    except OSError:
        # Windows access control is managed by the user's account and folder permissions.
        pass


def clear_api_key() -> None:
    if _PRIVATE_FILE.exists():
        _PRIVATE_FILE.unlink()


def _api_key() -> str | None:
    return settings.gemini_api_key or _read_private().get("gemini_api_key")


def _safe_document_path(document_id: str) -> Path:
    if not re.fullmatch(r"[0-9a-f-]{36}", document_id):
        raise ValueError("Invalid document identifier.")
    path = (_DOC_DIR / f"{document_id}.txt").resolve()
    root = _DOC_DIR.resolve()
    if root not in path.parents:
        raise ValueError("Invalid document path.")
    return path


def _extract_pdf(path: Path) -> tuple[str, int]:
    reader = PdfReader(str(path))
    chunks: list[str] = []
    for page in reader.pages[:300]:
        text = page.extract_text() or ""
        if text.strip():
            chunks.append(text)
        if sum(len(item) for item in chunks) >= _MAX_EXTRACTED_CHARS:
            break
    return "\n\n".join(chunks)[:_MAX_EXTRACTED_CHARS], len(reader.pages)


def _extract_docx(path: Path) -> tuple[str, int]:
    document = Document(str(path))
    chunks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            line = " | ".join(cell.text.strip() for cell in row.cells)
            if line.strip(" |"):
                chunks.append(line)
    return "\n".join(chunks)[:_MAX_EXTRACTED_CHARS], len(document.paragraphs)


def store_uploaded_document(path: Path, original_name: str) -> dict[str, Any]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text, pages = _extract_pdf(path)
        kind = "PDF"
    elif suffix == ".docx":
        text, pages = _extract_docx(path)
        kind = "Word document"
    else:
        raise ValueError("Only PDF and DOCX files are supported.")
    if not text.strip():
        raise ValueError("No readable text was found in this document.")
    _DOC_DIR.mkdir(parents=True, exist_ok=True)
    document_id = str(uuid.uuid4())
    target = _safe_document_path(document_id)
    header = json.dumps({"name": original_name[:180], "kind": kind, "pages": pages}, ensure_ascii=False)
    target.write_text(header + "\n" + text, encoding="utf-8")
    supabase_state.upload_runtime_file(target, namespace="assistant_documents")
    return {"document_id": document_id, "name": original_name, "kind": kind, "pages": pages, "characters": len(text)}


def attach_saved_report(report_id: str) -> dict[str, Any]:
    path = get_report(report_id)
    record = report_record(report_id)
    if not path or not path.exists() or not record:
        raise FileNotFoundError("Saved report was not found.")
    return store_uploaded_document(path, path.name)


def read_document(document_id: str) -> tuple[dict[str, Any], str]:
    path = _safe_document_path(document_id)
    if not path.exists():
        supabase_state.restore_runtime_file(path, namespace="assistant_documents")
    if not path.exists():
        raise FileNotFoundError("Attached document is no longer available.")
    raw = path.read_text(encoding="utf-8")
    first, _, text = raw.partition("\n")
    return json.loads(first), text


def _compact_context(context: dict[str, Any] | None) -> str:
    if not context:
        return ""
    # Limit size while retaining the most decision-relevant sections.
    selected = {
        key: context.get(key)
        for key in (
            "farm",
            "forecast_summary",
            "selected_frame",
            "hazards",
            "pest_risk",
            "pest_specific",
            "suitability",
            "farm_condition",
            "recommended_intervention",
            "rehabilitation_plan",
        )
        if context.get(key) is not None
    }
    return json.dumps(selected, ensure_ascii=False, default=str)[:28_000]


def _candidate_model_names(private: dict[str, Any]) -> list[str]:
    """Return stable candidates without exposing or persisting retired explicit model versions."""
    candidates: list[str] = []
    allowed = set(_MODEL_FALLBACKS) | {str(settings.gemini_model).strip().removeprefix("models/")}
    for candidate in (private.get("resolved_model"), settings.gemini_model, *_MODEL_FALLBACKS):
        value = str(candidate or "").strip().removeprefix("models/")
        if not value or value not in allowed or value in candidates:
            continue
        candidates.append(value)
    return candidates


def _provider_error_detail(response: httpx.Response | Any) -> str:
    try:
        payload = response.json()
    except Exception:
        return str(getattr(response, "text", "") or "")[:400]
    if isinstance(payload, dict):
        error = payload.get("error")
        if isinstance(error, dict):
            return str(error.get("message") or error.get("status") or "")[:500]
        feedback = payload.get("promptFeedback")
        if isinstance(feedback, dict) and feedback.get("blockReason"):
            return f"Request blocked: {feedback.get('blockReason')}"
    return ""


def _extract_answer(data: dict[str, Any]) -> tuple[str, str]:
    candidates = data.get("candidates") or []
    if not candidates:
        feedback = data.get("promptFeedback") or {}
        reason = feedback.get("blockReason") or "No response candidate was returned"
        return "", str(reason)
    candidate = candidates[0] or {}
    parts = (candidate.get("content") or {}).get("parts") or []
    answer = "\n".join(str(part.get("text", "")) for part in parts if isinstance(part, dict) and part.get("text")).strip()
    return answer, str(candidate.get("finishReason") or "")


async def _post_generation(client: httpx.AsyncClient, model_name: str, payload: dict[str, Any]) -> httpx.Response:
    endpoint = f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:generateContent"
    key = _api_key()
    assert key
    last_response: httpx.Response | None = None
    for attempt in range(3):
        try:
            response = await client.post(endpoint, headers={"x-goog-api-key": key}, json=payload)
        except httpx.TimeoutException as exc:
            if attempt == 2:
                raise RuntimeError("The AI provider timed out. Check your connection and try again.") from exc
            await __import__("asyncio").sleep(0.5 * (attempt + 1))
            continue
        except httpx.RequestError as exc:
            raise RuntimeError("The AI provider could not be reached. Check your internet connection and try again.") from exc
        last_response = response
        if response.status_code not in {500, 502, 503, 504}:
            return response
        if attempt < 2:
            await __import__("asyncio").sleep(0.6 * (attempt + 1))
    assert last_response is not None
    return last_response


async def chat_with_gemini(
    message: str,
    history: list[dict[str, str]],
    context: dict[str, Any] | None = None,
    document_ids: list[str] | None = None,
) -> dict[str, Any]:
    key = _api_key()
    if not key:
        raise RuntimeError("CoCO-PILOT is not configured. Add a Google AI API key in Settings.")

    context_parts: list[str] = []
    compact = _compact_context(context)
    if compact:
        context_parts.append("COCO-AID CURRENT CONTEXT:\n" + compact)

    attached: list[dict[str, Any]] = []
    remaining_document_chars = 48_000
    for document_id in (document_ids or [])[:3]:
        meta, text = read_document(document_id)
        attached.append(meta)
        excerpt = text[:min(20_000, remaining_document_chars)]
        remaining_document_chars -= len(excerpt)
        if excerpt:
            context_parts.append(f"ATTACHED {meta.get('kind')} - {meta.get('name')}:\n{excerpt}")
        if remaining_document_chars <= 0:
            break

    user_text = message.strip()
    if context_parts:
        user_text = "\n\n".join(context_parts) + "\n\nUSER QUESTION:\n" + user_text

    contents: list[dict[str, Any]] = []
    for item in history[-8:]:
        role = "model" if item.get("role") == "assistant" else "user"
        text = str(item.get("content", "")).strip()
        if text:
            contents.append({"role": role, "parts": [{"text": text[:6_000]}]})
    contents.append({"role": "user", "parts": [{"text": user_text[:85_000]}]})

    base_payload = {
        "system_instruction": {"parts": [{"text": SYSTEM_PROMPT}]},
        "contents": contents,
        "generationConfig": {"temperature": 0.25, "maxOutputTokens": 1800, "topP": 0.9},
    }
    private = _read_private()
    models = _candidate_model_names(private)
    if not models:
        models = ["gemini-flash-latest"]

    selected_model: str | None = None
    answer = ""
    finish_reason = ""
    last_detail = ""

    async with httpx.AsyncClient(timeout=httpx.Timeout(75.0, connect=15.0)) as client:
        for model_name in models:
            response = await _post_generation(client, model_name, base_payload)
            if response.status_code < 400:
                try:
                    data = response.json()
                except ValueError as exc:
                    last_detail = "The AI provider returned an unreadable response."
                    continue
                answer, finish_reason = _extract_answer(data)
                if answer:
                    selected_model = model_name
                    break
                last_detail = finish_reason or "The AI provider returned no readable text."
                continue

            detail = _provider_error_detail(response)
            last_detail = detail or f"Provider status {response.status_code}"
            if response.status_code == 429:
                raise RuntimeError("The AI free-tier limit was reached. Wait briefly, then try again.")
            detail_lower = last_detail.lower()
            incompatible = response.status_code in {400, 404} and any(
                marker in detail_lower for marker in ("not found", "not available", "unsupported", "model")
            )
            if incompatible or response.status_code >= 500:
                continue
            if response.status_code in {401, 403}:
                raise RuntimeError("The API key was rejected. Check the key in Settings and try again.")
            raise RuntimeError(f"The AI request was rejected: {last_detail}")

        if selected_model and answer and finish_reason.upper() in {"MAX_TOKENS", "LENGTH"}:
            continuation_contents = list(contents)
            continuation_contents.append({"role": "model", "parts": [{"text": answer}]})
            continuation_contents.append({
                "role": "user",
                "parts": [{"text": "Continue exactly where the response stopped. Finish the answer concisely without repeating earlier text."}],
            })
            continuation_payload = {
                "system_instruction": {"parts": [{"text": SYSTEM_PROMPT}]},
                "contents": continuation_contents,
                "generationConfig": {"temperature": 0.2, "maxOutputTokens": 1000, "topP": 0.9},
            }
            continuation_response = await _post_generation(client, selected_model, continuation_payload)
            if continuation_response.status_code < 400:
                try:
                    continuation, _ = _extract_answer(continuation_response.json())
                except ValueError:
                    continuation = ""
                if continuation:
                    answer = (answer.rstrip() + "\n" + continuation.lstrip())[:12_000]

    if not selected_model or not answer:
        raise RuntimeError(f"CoCO-PILOT could not get a usable response. {last_detail or 'Try again shortly.'}")

    if selected_model != private.get("resolved_model") and not settings.gemini_api_key:
        updated = _read_private()
        if updated.get("gemini_api_key"):
            updated["resolved_model"] = selected_model
            updated["updated_at"] = datetime.now(UTC).isoformat()
            _PRIVATE_FILE.parent.mkdir(parents=True, exist_ok=True)
            _PRIVATE_FILE.write_text(json.dumps(updated, indent=2), encoding="utf-8")

    percentages = [
        float(value)
        for value in re.findall(r"(?<!\d)(\d{1,3}(?:\.\d+)?)\s*%", answer)
        if 0 <= float(value) <= 100
    ][:4]
    return {
        "answer": answer,
        "percentages": percentages,
        "model": "Automatic compatible Flash model",
        "attached_documents": attached,
        "complete": finish_reason.upper() not in {"MAX_TOKENS", "LENGTH"} or len(answer) > 0,
    }
