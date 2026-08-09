from __future__ import annotations

import hashlib
import json
import os
from datetime import UTC
from io import BytesIO
from pathlib import Path
from typing import Any
from uuid import UUID

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont, TTFError
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from app.coco_pilot.context import build_context
from app.coco_pilot import repository
from app.core.config import settings
from app.domain.coco_pilot import FormalReportRecord, FormalReportRequest

FORMAL_REPORT_GENERATOR_VERSION = "formal-report-generator-1.1.0"
FORMAL_REPORT_DATA_NOTICE = (
    "This report is generated from a saved COCOAID integrated decision-support record. "
    "Numeric tables are populated directly from versioned analytical outputs. Narrative text does not "
    "override the source engines, create new field evidence, authorize unverified chemical dosage, or guarantee income or recovery."
)

def _font_candidates() -> dict[str, list[Path]]:
    windir = Path(os.environ.get("WINDIR", "C:/Windows"))
    return {
        "regular": [
            windir / "Fonts" / "times.ttf",
            Path("/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman.ttf"),
            Path("/usr/share/fonts/truetype/tinos/Tinos-Regular.ttf"),
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf"),
        ],
        "bold": [
            windir / "Fonts" / "timesbd.ttf",
            Path("/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman_Bold.ttf"),
            Path("/usr/share/fonts/truetype/tinos/Tinos-Bold.ttf"),
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSerif-Bold.ttf"),
        ],
        "italic": [
            windir / "Fonts" / "timesi.ttf",
            Path("/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman_Italic.ttf"),
            Path("/usr/share/fonts/truetype/tinos/Tinos-Italic.ttf"),
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSerif-Italic.ttf"),
        ],
        "bold_italic": [
            windir / "Fonts" / "timesbi.ttf",
            Path("/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman_Bold_Italic.ttf"),
            Path("/usr/share/fonts/truetype/tinos/Tinos-BoldItalic.ttf"),
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSerif-BoldItalic.ttf"),
        ],
    }


def _first_report_font_file(paths: list[Path]) -> Path | None:
    """Return a real font file, never a directory or placeholder path."""
    for path in paths:
        try:
            if path.is_file():
                return path
        except OSError:
            continue
    return None


def _register_report_fonts() -> tuple[str, str, str, str]:
    names = {
        "regular": "COCOAID-Times",
        "bold": "COCOAID-Times-Bold",
        "italic": "COCOAID-Times-Italic",
        "bold_italic": "COCOAID-Times-BoldItalic",
    }
    resolved = {style: _first_report_font_file(paths) for style, paths in _font_candidates().items()}
    if not all(resolved.values()):
        return "Times-Roman", "Times-Bold", "Times-Italic", "Times-BoldItalic"

    try:
        for style, path in resolved.items():
            if names[style] not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(names[style], str(path)))
        pdfmetrics.registerFontFamily(
            "COCOAID-Times",
            normal=names["regular"],
            bold=names["bold"],
            italic=names["italic"],
            boldItalic=names["bold_italic"],
        )
    except (OSError, TTFError, ValueError):
        # Report fonts are optional. Render/CI images may not ship a Times-compatible
        # TTF family, so use ReportLab's built-in Times metrics instead of crashing.
        return "Times-Roman", "Times-Bold", "Times-Italic", "Times-BoldItalic"

    return names["regular"], names["bold"], names["italic"], names["bold_italic"]


PDF_FONT_REGULAR, PDF_FONT_BOLD, PDF_FONT_ITALIC, PDF_FONT_BOLD_ITALIC = _register_report_fonts()



def _canonical(value: Any) -> str:
    return json.dumps(value, sort_keys=True, separators=(",", ":"), ensure_ascii=False, default=str, allow_nan=False)


def _sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _fmt(value: Any, digits: int = 2) -> str:
    if value is None:
        return "Not available"
    if isinstance(value, bool):
        return "Yes" if value else "No"
    try:
        return f"{float(value):,.{digits}f}"
    except (TypeError, ValueError):
        return str(value)


def _pct(value: Any) -> str:
    return "Not available" if value is None else f"{float(value):.1%}"


def _component(decision: dict[str, Any], name: str) -> dict[str, Any]:
    return next((item for item in decision.get("component_results", []) if item.get("component") == name), {})


def _set_cell_text(cell, text: Any, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(9)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _docx_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)
    for name, size in (("Title", 18), ("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)):
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.color.rgb = None


def _add_docx_table(document: Document, headers: list[str], rows: list[list[Any]], widths: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[index], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            _set_cell_text(cells[index], value)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    document.add_paragraph()


def _add_docx_paragraph(document: Document, text: str, *, bold: bool = False, italic: bool = False, center: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic



def _add_docx_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(9)


def _configure_docx_office_header(section, analysis_run_id: Any) -> None:
    header = section.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("COCOAID | INTEGRATED DECISION-SUPPORT REPORT")
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(9)
    run.bold = True
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(f"Analysis Record {analysis_run_id} | Page ")
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(9)
    _add_docx_page_number(p)


def _docx_report(
    decision: dict[str, Any], narrative: dict[str, Any] | None, source_manifest: list[dict[str, Any]],
    linked_records: dict[str, Any], request: FormalReportRequest,
) -> bytes:
    document = Document()
    _docx_styles(document)
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    _configure_docx_office_header(section, decision.get("analysis_run_id"))

    title = request.title or "COCOAID Integrated Decision-Support Report"
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(18)
    _add_docx_paragraph(
        document,
        "Bayesian Probabilistic Agroecosystem Simulation and Geospatial AI-Based Decision-Support Framework",
        center=True,
    )
    _add_docx_paragraph(document, "OFFICIAL RESEARCH DECISION-SUPPORT RECORD", bold=True, center=True)
    _add_docx_table(document, ["Document Control", "Value"], [
        ["Analysis run", decision.get("analysis_run_id")],
        ["Generated at", decision.get("generated_at")],
        ["Classification", "Research prototype analytical record"],
        ["Generator", FORMAL_REPORT_GENERATOR_VERSION],
    ], widths=[2.2, 4.4])
    _add_docx_paragraph(document, FORMAL_REPORT_DATA_NOTICE, bold=True)

    overview = decision.get("overview", {})
    summary = decision.get("summary", {})
    document.add_heading("1. Executive Summary", level=1)
    if narrative:
        _add_docx_paragraph(document, narrative.get("conclusion") or narrative.get("full_text", ""))
    _add_docx_table(document, ["Field", "Value"], [
        ["Analysis run", decision.get("analysis_run_id")],
        ["Status", decision.get("status")],
        ["Generated at", decision.get("generated_at")],
        ["Production estimate", f"{_fmt(overview.get('production_estimate'), 3)} {overview.get('production_unit', '')}"],
        ["Production interval", f"{_fmt(overview.get('production_lower'), 3)} to {_fmt(overview.get('production_upper'), 3)}"],
        ["Probability of decline", _pct(overview.get("probability_of_decline"))],
        ["Probability of recovery", _pct(overview.get("probability_of_recovery"))],
        ["Data completeness", _pct(overview.get("data_completeness"))],
        ["Urgent recommendations", overview.get("urgent_recommendation_count", 0)],
    ], widths=[2.2, 4.4])

    document.add_heading("2. Analytical Component Status", level=1)
    component_rows = []
    for item in decision.get("component_results", []):
        component_rows.append([
            item.get("component"), item.get("engine_id"), item.get("status"),
            item.get("record_id") or "Not attached", "; ".join(item.get("warnings", []) or item.get("errors", [])) or "None",
        ])
    _add_docx_table(document, ["Component", "Engine", "Status", "Record ID", "Disclosure"], component_rows)

    document.add_heading("3. Production and Bayesian Uncertainty", level=1)
    prod = _component(decision, "production").get("summary", {})
    bayes = _component(decision, "bayesian").get("summary", {})
    distribution = bayes.get("production_distribution") or {}
    _add_docx_table(document, ["Metric", "Value"], [
        ["Raw ML prediction", f"{_fmt(prod.get('raw_ml_prediction'), 3)} {prod.get('unit', '')}"],
        ["Variety-adjusted prediction", f"{_fmt(prod.get('variety_adjusted_prediction'), 3)} {prod.get('unit', '')}"],
        ["Posterior lower", _fmt(distribution.get("lower"), 3)],
        ["Posterior median", _fmt(distribution.get("median"), 3)],
        ["Posterior upper", _fmt(distribution.get("upper"), 3)],
        ["Decline probability", _pct(bayes.get("probability_of_decline"))],
        ["Recovery probability", _pct(bayes.get("probability_of_recovery"))],
        ["Tree mortality probability", _pct(bayes.get("probability_of_tree_mortality"))],
        ["Pest outbreak probability", _pct(bayes.get("probability_of_pest_outbreak"))],
    ])

    document.add_heading("4. Pest, Intercropping, and Rehabilitation Overview", level=1)
    pest = _component(decision, "pest").get("summary", {})
    intercrop = _component(decision, "intercropping").get("summary", {})
    rehabilitation = _component(decision, "rehabilitation").get("summary", {})
    highest = pest.get("highest_risk") or {}
    best = intercrop.get("best_candidate") or {}
    selected = rehabilitation.get("selected_scenario_result") or {}
    _add_docx_table(document, ["Domain", "Primary result", "Supporting value"], [
        ["Pest", highest.get("pest_profile_id") or "Not available", _pct(highest.get("outbreak_probability"))],
        ["Intercropping", best.get("candidate_id") or "Not available", f"{_fmt(best.get('suitability_score'), 1)}/100"],
        ["Rehabilitation", rehabilitation.get("selected_scenario") or "Not available", f"PHP {_fmt(selected.get('total_cost_php'), 2)}"],
        ["Scenario severe-loss probability", "Selected scenario", _pct(selected.get("severe_loss_probability"))],
        ["Scenario labor", "Selected scenario", f"{_fmt(selected.get('labor_person_days'), 1)} person-days"],
    ])


    rehabilitation_plan = linked_records.get("rehabilitation_plan") or {}
    scenarios = list(rehabilitation_plan.get("scenarios") or [])
    if scenarios:
        document.add_heading("5. Scenario Comparison", level=1)
        scenario_order = {"no_action": 0, "pest_management": 1, "fertilization": 2, "replanting": 3, "intercropping": 4, "combined_rehabilitation": 5}
        scenarios = sorted(scenarios, key=lambda item: scenario_order.get(str(item.get("scenario_type")), 99))
        _add_docx_table(document, ["Scenario", "Status", "Cost (PHP)", "Labor", "Severe-loss probability", "Expected utility"], [[
            str(item.get("scenario_type") or "unknown").replace("_", " ").title(),
            item.get("status"), _fmt(item.get("total_cost_php"), 2),
            f"{_fmt(item.get('labor_person_days'), 1)} person-days",
            _pct(item.get("severe_loss_probability")), _fmt(item.get("expected_utility"), 3),
        ] for item in scenarios])

        selected_record = next((item for item in scenarios if item.get("scenario_type") == rehabilitation_plan.get("selected_scenario")), {})
        selected_ids = {str(item) for item in selected_record.get("action_ids", [])}
        selected_actions = [item for item in rehabilitation_plan.get("actions", []) if str(item.get("id") or item.get("action_id")) in selected_ids]
        if selected_actions:
            document.add_heading("6. Selected Scenario Work Plan", level=1)
            _add_docx_table(document, ["Action", "Problem", "Schedule", "Cost (PHP)", "Field confirmation"], [[
                str(item.get("action_type") or "action").replace("_", " ").title(),
                item.get("problem_detected"), item.get("scheduled_date") or "After verification",
                _fmt(item.get("total_php"), 2), "Yes" if item.get("requires_field_confirmation") else "No",
            ] for item in selected_actions])

    document.add_heading("7. Prioritized Recommendations", level=1)
    recommendation_rows = []
    for item in decision.get("recommendations", []):
        recommendation_rows.append([
            item.get("priority"), item.get("category"), item.get("title"), item.get("action"),
            "Yes" if item.get("requires_field_confirmation") else "No",
        ])
    _add_docx_table(document, ["Priority", "Category", "Recommendation", "Action", "Field confirmation"], recommendation_rows)

    document.add_heading("8. Evidence Traceability", level=1)
    evidence_rows: list[list[Any]] = []
    for item in decision.get("recommendations", []):
        for evidence in item.get("evidence", []):
            evidence_rows.append([
                item.get("title"), evidence.get("source_component"), evidence.get("record_id"),
                evidence.get("field"), str(evidence.get("value")), evidence.get("explanation"),
            ])
    _add_docx_table(document, ["Recommendation", "Source", "Record", "Field", "Value", "Meaning"], evidence_rows)

    document.add_heading("9. Source Provenance", level=1)
    source_rows = [[
        item.get("source_type"), item.get("title"), item.get("source_id"),
        item.get("organization") or item.get("engine_id") or "Not specified",
        item.get("access_class", "analytical"),
    ] for item in source_manifest]
    _add_docx_table(document, ["Type", "Title", "Source ID", "Organization or engine", "Access class"], source_rows)

    document.add_heading("10. Limitations and Safety Boundaries", level=1)
    limitations = list((decision.get("provenance") or {}).get("limitations", []))
    if narrative:
        limitations.extend(narrative.get("limitations", []))
    limitations.extend([
        "Predicted or suspected hazards are not treated as confirmed field damage.",
        "No chemical dosage is generated by the formal-report layer.",
        "Farmer names, protected identities, and restricted raw records are excluded.",
    ])
    for item in dict.fromkeys(limitations):
        p = document.add_paragraph(style="List Bullet")
        run = p.add_run(str(item))
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(11)

    document.add_heading("11. Reproducibility Record", level=1)
    _add_docx_table(document, ["Field", "Value"], [
        ["Decision-support parameter version", decision.get("parameter_version")],
        ["Dependency-graph version", decision.get("dependency_graph_version")],
        ["Report-generator version", FORMAL_REPORT_GENERATOR_VERSION],
        ["Farm data version", (decision.get("provenance") or {}).get("farm_data_version")],
        ["Weather run", (decision.get("provenance") or {}).get("weather_run_id")],
    ])

    core = document.core_properties
    core.title = title
    core.author = "COCOAID Research Prototype"
    core.subject = "Versioned integrated decision-support report"
    core.comments = FORMAL_REPORT_DATA_NOTICE[:255]
    core.created = request.generated_at.astimezone(UTC).replace(tzinfo=None)
    core.modified = request.generated_at.astimezone(UTC).replace(tzinfo=None)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="TNRBody", parent=styles["BodyText"], fontName=PDF_FONT_REGULAR, fontSize=10, leading=13, alignment=TA_JUSTIFY, textColor=colors.black))
    styles.add(ParagraphStyle(name="TNRTitle", parent=styles["Title"], fontName=PDF_FONT_BOLD, fontSize=18, leading=22, alignment=TA_CENTER, textColor=colors.black))
    styles.add(ParagraphStyle(name="TNRH1", parent=styles["Heading1"], fontName=PDF_FONT_BOLD, fontSize=13, leading=16, textColor=colors.black, spaceBefore=8, spaceAfter=5))
    return styles


def _pdf_table(headers: list[str], rows: list[list[Any]], widths=None) -> Table:
    data = [[Paragraph(str(value), ParagraphStyle(name=f"h{id(value)}", fontName=PDF_FONT_BOLD, fontSize=8, leading=10)) for value in headers]]
    body_style = ParagraphStyle(name="tbody", fontName=PDF_FONT_REGULAR, fontSize=7.5, leading=9.5)
    for row in rows:
        data.append([Paragraph(str(value), body_style) for value in row])
    table = Table(data, colWidths=widths, repeatRows=1, hAlign="LEFT")
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.35, colors.black),
        ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 3),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    return table



def _pdf_office_header_footer(canvas, doc) -> None:
    canvas.saveState()
    width, height = A4
    canvas.setStrokeColor(colors.HexColor("#777777"))
    canvas.setLineWidth(0.35)
    canvas.line(16 * mm, height - 11 * mm, width - 16 * mm, height - 11 * mm)
    canvas.setFont(PDF_FONT_BOLD, 7.5)
    canvas.setFillColor(colors.black)
    canvas.drawString(16 * mm, height - 8.5 * mm, "COCOAID | INTEGRATED DECISION-SUPPORT REPORT")
    canvas.line(16 * mm, 11 * mm, width - 16 * mm, 11 * mm)
    canvas.setFont(PDF_FONT_REGULAR, 7.5)
    canvas.drawString(16 * mm, 7 * mm, "Research prototype analytical record")
    canvas.drawRightString(width - 16 * mm, 7 * mm, f"Page {doc.page}")
    canvas.restoreState()


def _pdf_report(
    decision: dict[str, Any], narrative: dict[str, Any] | None, source_manifest: list[dict[str, Any]],
    linked_records: dict[str, Any], request: FormalReportRequest,
) -> bytes:
    buffer = BytesIO()
    styles = _pdf_styles()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=16 * mm, rightMargin=16 * mm, topMargin=18 * mm, bottomMargin=17 * mm)
    story: list[Any] = [
        Spacer(1, 20 * mm),
        Paragraph(request.title or "COCOAID Integrated Decision-Support Report", styles["TNRTitle"]),
        Paragraph("Bayesian Probabilistic Agroecosystem Simulation and Geospatial AI-Based Decision-Support Framework", styles["TNRBody"]),
        Spacer(1, 6),
        Paragraph("OFFICIAL RESEARCH DECISION-SUPPORT RECORD", styles["TNRH1"]),
        Spacer(1, 8),
        _pdf_table(["Document Control", "Value"], [
            ["Analysis run", decision.get("analysis_run_id")],
            ["Generated at", decision.get("generated_at")],
            ["Classification", "Research prototype analytical record"],
            ["Generator", FORMAL_REPORT_GENERATOR_VERSION],
        ], [52 * mm, 120 * mm]),
        Spacer(1, 8), Paragraph(FORMAL_REPORT_DATA_NOTICE, styles["TNRBody"]), PageBreak(),
    ]
    overview = decision.get("overview", {})
    story.extend([
        Paragraph("1. Executive Summary", styles["TNRH1"]),
        Paragraph((narrative or {}).get("conclusion") or "The report summarizes the saved integrated analytical record.", styles["TNRBody"]),
        _pdf_table(["Field", "Value"], [
            ["Analysis run", decision.get("analysis_run_id")],
            ["Status", decision.get("status")],
            ["Production estimate", f"{_fmt(overview.get('production_estimate'), 3)} {overview.get('production_unit', '')}"],
            ["Probability of decline", _pct(overview.get("probability_of_decline"))],
            ["Probability of recovery", _pct(overview.get("probability_of_recovery"))],
            ["Data completeness", _pct(overview.get("data_completeness"))],
        ], [52 * mm, 120 * mm]), Spacer(1, 8),
        Paragraph("2. Component Status", styles["TNRH1"]),
        _pdf_table(["Component", "Engine", "Status", "Record ID"], [[
            item.get("component"), item.get("engine_id"), item.get("status"), item.get("record_id") or "Not attached"
        ] for item in decision.get("component_results", [])], [28 * mm, 42 * mm, 24 * mm, 78 * mm]),
        Spacer(1, 8), Paragraph("3. Prioritized Recommendations", styles["TNRH1"]),
        _pdf_table(["Priority", "Category", "Recommendation", "Action"], [[
            item.get("priority"), item.get("category"), item.get("title"), item.get("action")
        ] for item in decision.get("recommendations", [])], [20 * mm, 24 * mm, 50 * mm, 78 * mm]),
    ])
    rehabilitation_plan = linked_records.get("rehabilitation_plan") or {}
    scenarios = list(rehabilitation_plan.get("scenarios") or [])
    if scenarios:
        scenario_order = {"no_action": 0, "pest_management": 1, "fertilization": 2, "replanting": 3, "intercropping": 4, "combined_rehabilitation": 5}
        scenarios = sorted(scenarios, key=lambda item: scenario_order.get(str(item.get("scenario_type")), 99))
        story.extend([
            Spacer(1, 8), Paragraph("4. Scenario Comparison", styles["TNRH1"]),
            _pdf_table(["Scenario", "Status", "Cost", "Labor", "Loss risk", "Utility"], [[
                str(item.get("scenario_type") or "unknown").replace("_", " ").title(), item.get("status"),
                f"PHP {_fmt(item.get('total_cost_php'), 2)}", _fmt(item.get("labor_person_days"), 1),
                _pct(item.get("severe_loss_probability")), _fmt(item.get("expected_utility"), 3),
            ] for item in scenarios], [38 * mm, 20 * mm, 30 * mm, 20 * mm, 28 * mm, 24 * mm]),
        ])
    story.extend([PageBreak(), Paragraph("5. Evidence Traceability", styles["TNRH1"])])
    evidence_rows = []
    for item in decision.get("recommendations", []):
        for evidence in item.get("evidence", []):
            evidence_rows.append([item.get("title"), evidence.get("source_component"), evidence.get("field"), str(evidence.get("value")), evidence.get("explanation")])
    story.extend([
        _pdf_table(["Recommendation", "Source", "Field", "Value", "Meaning"], evidence_rows, [42 * mm, 24 * mm, 30 * mm, 18 * mm, 58 * mm]),
        Spacer(1, 8), Paragraph("6. Source Provenance", styles["TNRH1"]),
        _pdf_table(["Type", "Title", "Source ID", "Access"], [[
            item.get("source_type"), item.get("title"), item.get("source_id"), item.get("access_class", "analytical")
        ] for item in source_manifest], [28 * mm, 70 * mm, 54 * mm, 20 * mm]),
        Spacer(1, 8), Paragraph("7. Limitations", styles["TNRH1"]),
    ])
    for item in dict.fromkeys((decision.get("provenance") or {}).get("limitations", []) + (narrative or {}).get("limitations", []) + [
        "Predicted or suspected hazards are not confirmed field damage.",
        "No unverified chemical dosage is generated.",
        "Protected farmer identities and restricted raw records are excluded.",
    ]):
        story.append(Paragraph(f"- {item}", styles["TNRBody"]))
    doc.build(story, onFirstPage=_pdf_office_header_footer, onLaterPages=_pdf_office_header_footer)
    return buffer.getvalue()


def generate_formal_report(
    request: FormalReportRequest, *, database_path: Path | None = None,
) -> tuple[FormalReportRecord, Path]:
    context = build_context(request.analysis_run_id, database_path=database_path)
    decision = context["decision"]
    linked_records = context["linked_records"]
    narrative = repository.get_response(request.narrative_run_id, database_path=database_path) if request.narrative_run_id else None
    if request.narrative_run_id and not narrative:
        raise FileNotFoundError("CoCO-PILOT narrative run was not found.")
    if narrative and narrative.get("analysis_run_id") != str(request.analysis_run_id):
        raise ValueError("Narrative run does not belong to the requested decision-support run.")

    source_manifest = context["source_manifest"]
    fingerprint_payload = {
        "decision": decision,
        "narrative": narrative,
        "source_manifest": source_manifest,
        "linked_records": linked_records,
        "generator_version": FORMAL_REPORT_GENERATOR_VERSION,
        "report_format": request.report_format,
        "title": request.title,
    }
    content_fingerprint = hashlib.sha256(_canonical(fingerprint_payload).encode("utf-8")).hexdigest()
    data = _docx_report(decision, narrative, source_manifest, linked_records, request) if request.report_format == "docx" else _pdf_report(decision, narrative, source_manifest, linked_records, request)
    file_sha256 = _sha256_bytes(data)

    settings.reports_dir.mkdir(parents=True, exist_ok=True)
    extension = request.report_format
    filename = f"COCOAID_Formal_Report_{str(request.analysis_run_id)[:8]}_{content_fingerprint[:8]}.{extension}"
    path = settings.reports_dir / filename
    path.write_bytes(data)
    record = FormalReportRecord(
        analysis_run_id=request.analysis_run_id,
        narrative_run_id=request.narrative_run_id,
        report_format=request.report_format,
        filename=filename,
        file_sha256=file_sha256,
        content_fingerprint=content_fingerprint,
        generator_version=FORMAL_REPORT_GENERATOR_VERSION,
        source_manifest=source_manifest,
        warnings=list(decision.get("warnings", []))[:50],
        data_notice=FORMAL_REPORT_DATA_NOTICE,
        created_at=request.generated_at.astimezone(UTC),
    )
    repository.save_report(record, path, database_path=database_path)
    return record, path
