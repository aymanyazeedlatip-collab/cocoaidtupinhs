from __future__ import annotations

from datetime import UTC, datetime
from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient
from pypdf import PdfReader

from app.engines.decision_support import DecisionSupportEngine
from app.main import app
from tests.phase9_factory import decision_request, prepare_phase9_records

client = TestClient(app)


def _decision_run_id():
    production, posterior, pest, intercrop, rehabilitation = prepare_phase9_records()
    output = DecisionSupportEngine().execute(
        decision_request(production, posterior, pest, intercrop, rehabilitation)
    ).output
    return output.record.analysis_run_id


def test_phase10_status_explain_report_and_download():
    run_id = _decision_run_id()
    status = client.get("/api/v2/coco-pilot/status")
    assert status.status_code == 200
    assert status.json()["availability"] == "available"
    assert status.json()["safety_policy"]["numeric_tables_generated_by_llm"] is False

    explanation = client.post("/api/v2/coco-pilot/explain", json={
        "analysis_run_id": str(run_id),
        "mode": "report_narrative",
        "provider_mode": "deterministic",
        "include_pca_references": True,
        "generated_at": datetime(2026, 8, 4, 6, tzinfo=UTC).isoformat(),
    })
    assert explanation.status_code == 200, explanation.text
    narrative = explanation.json()
    assert narrative["provider"] == "deterministic"
    assert narrative["citations"]
    assert narrative["redaction_summary"]["farmer_names_included"] is False

    report_ids = []
    for report_format in ("docx", "pdf"):
        generated = client.post("/api/v2/formal-reports/generate", json={
            "analysis_run_id": str(run_id),
            "narrative_run_id": narrative["run_id"],
            "report_format": report_format,
            "generated_at": datetime(2026, 8, 4, 7, tzinfo=UTC).isoformat(),
        })
        assert generated.status_code == 200, generated.text
        item = generated.json()
        report_ids.append(item["report_id"])
        assert len(item["file_sha256"]) == 64
        download = client.get(item["download_url"])
        assert download.status_code == 200
        if report_format == "docx":
            document = Document(BytesIO(download.content))
            text = "\n".join(p.text for p in document.paragraphs)
            assert "Integrated Decision-Support Report" in text
            assert "OFFICIAL RESEARCH DECISION-SUPPORT RECORD" in text
            assert "Scenario Comparison" in text
            assert "Prioritized Recommendations" in text
            assert "Private Farmer" not in text
            assert document.styles["Normal"].font.name == "Times New Roman"
            assert "COCOAID | INTEGRATED DECISION-SUPPORT REPORT" in document.sections[0].header.paragraphs[0].text
            assert "Analysis Record" in document.sections[0].footer.paragraphs[0].text
        else:
            reader = PdfReader(BytesIO(download.content))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            assert "Integrated Decision-Support Report" in text
            assert "OFFICIAL RESEARCH DECISION-SUPPORT RECORD" in text
            assert "Evidence Traceability" in text
            assert "Research prototype analytical record" in text

    listing = client.get(f"/api/v2/formal-reports?analysis_run_id={run_id}")
    assert listing.status_code == 200
    assert listing.json()["count"] == 2
    record = client.get(f"/api/v2/formal-reports/{report_ids[0]}")
    assert record.status_code == 200
    assert "filepath" not in record.json()


def test_phase10_health_contract_and_migration():
    health = client.get("/api/v2/health")
    assert health.status_code == 200
    assert health.json()["contract_api_version"] == "3.0.0-draft.10"
    migrations = health.json()["database_migrations"]
    assert migrations[9]["name"] == "phase10_coco_pilot_formal_reports"
    assert migrations[9]["state"] == "applied"
